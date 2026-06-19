"""Import Low Carb .docx recipes into recipes.json / recipes.js."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document

ING_MARKERS = re.compile(
    r"^(ingredients?|crust|marinade|for the .+:|thai tea mix directions)\s*:?\s*$",
    re.I,
)
DIR_MARKERS = re.compile(
    r"^(directions?|instructions?|how to prepare|method)\s*:?\s*$",
    re.I,
)
SKIP_LINES = re.compile(
    r"^(makes \d|oven:|effort:|\*+|note[s]?\s*:?\s*$)",
    re.I,
)


def slugify(text: str) -> str:
    s = text.lower()
    s = re.sub(r"[^a-z0-9]+", "-", s).strip("-")
    return f"low-carb-{s[:50]}"


def guess_category(title: str) -> str:
    t = title.lower()
    if any(w in t for w in ("cookie", "cake", "brownie", "cheesecake", "cupcake", "muffin", "frosting", "pudding", "ice cream", "shortbread")):
        return "dessert"
    if any(w in t for w in ("sauce", "tea", "blend", "dressing")):
        return "other"
    if "salad" in t:
        return "salad"
    return "main"


def parse_docx(path: Path) -> dict:
    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    if not lines:
        title = path.stem
        return {
            "id": slugify(title),
            "title": title,
            "collection": "low-carb",
            "category": guess_category(title),
            "ingredients": [],
            "steps": ["Recipe file was empty."],
            "pending": False,
        }

    title = lines[0]
    if ING_MARKERS.match(title) or len(title) > 60:
        title = path.stem.replace("-", " ")

    ingredients: list[str] = []
    steps: list[str] = []
    notes: list[str] = []
    section = "pre"
    step_buf: list[str] = []

    def flush_step() -> None:
        if step_buf:
            steps.append(" ".join(step_buf))
            step_buf.clear()

    for line in lines[1:]:
        if ING_MARKERS.match(line):
            flush_step()
            section = "ing"
            continue
        if DIR_MARKERS.match(line):
            flush_step()
            section = "dir"
            continue
        if line.startswith("*") and section == "pre":
            notes.append(line.lstrip("* ").strip())
            continue
        if SKIP_LINES.match(line) and section != "dir":
            if "oven" in line.lower():
                ingredients.append(line)
            continue

        if section == "ing":
            if re.match(r"^\d+[\.)]\s", line) and not ingredients:
                section = "dir"
                step_buf.append(line)
            else:
                ingredients.append(line)
        elif section == "dir":
            if re.match(r"^\d+[\.)]\s", line):
                flush_step()
                step_buf.append(re.sub(r"^\d+[\.)]\s*", "", line))
            elif step_buf and not line[0].isdigit():
                step_buf.append(line)
            else:
                flush_step()
                steps.append(line)
        else:
            if line.endswith(":") and len(line) < 40:
                ingredients.append(line)
                section = "ing"
            elif any(k in line.lower() for k in ("preheat", "bake", "mix", "combine", "stir", "pour", "cook", "blend", "refrigerate")):
                section = "dir"
                steps.append(line)
            elif len(line) > 80 or ";" in line:
                section = "dir"
                steps.append(line)
            else:
                ingredients.append(line)

    flush_step()

    if not ingredients and not steps:
        body = lines[1:] if lines[0] == title else lines
        mid = max(1, len(body) // 2)
        ingredients = body[:mid]
        steps = body[mid:] or ["See original docx in recipe hole folder."]

    return {
        "id": slugify(title),
        "title": title,
        "collection": "low-carb",
        "category": guess_category(title),
        "ingredients": ingredients,
        "steps": steps,
        "notes": notes,
        "pending": False,
    }


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    docx_dir = root / "recipe hole" / "Low Carb"
    data_path = root / "data" / "recipes.json"
    data = json.loads(data_path.read_text(encoding="utf-8"))

    recipes = []
    for path in sorted(docx_dir.glob("*.docx")):
        recipes.append(parse_docx(path))

    data["recipes"] = [r for r in data["recipes"] if r.get("collection") != "low-carb"]
    data["recipes"].extend(recipes)

    data["collections"] = [c for c in data["collections"] if c["id"] != "low-carb"]
    data["collections"].append(
        {
            "id": "low-carb",
            "title": "Low Carb",
            "subtitle": "Nice to have",
            "type": "text",
            "recipeCount": len(recipes),
        }
    )

    for coll in data["collections"]:
        if coll["id"] in ("prescott-cousins", "cabin"):
            coll["recipeCount"] = sum(
                1 for r in data["recipes"] if r.get("collection") == coll["id"]
            )

    payload = json.dumps(data, indent=2, ensure_ascii=False)
    data_path.write_text(payload + "\n", encoding="utf-8")
    (root / "data" / "recipes.js").write_text(
        f"window.RECIPE_DATA = {payload};\n", encoding="utf-8"
    )
    print(f"Imported {len(recipes)} low-carb docx recipes.")


if __name__ == "__main__":
    main()